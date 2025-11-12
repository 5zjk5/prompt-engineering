import requests
import re
import random
import time
from lxml import etree
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document


headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}


def get_title_url():
    """获得所有章节标题及 url"""
    url = 'https://www.doupocangqiong.org/doupocangqiong/'
    response = requests.get(url, headers=headers)
    response.encoding = 'utf8'
    html = etree.HTML(response.text)
    title_list = html.xpath('//*[@id="play_0"]/ul/li')
    titles = []
    links = []
    for title in title_list:
        t = title.xpath('./a/text()')[0]
        l = title.xpath('./a/@href')[0]
        l = 'https://www.doupocangqiong.org' + l
        titles.append(t)
        links.append(l)
    return titles, links


def get_text(url):
    """爬取当前章节正文内容"""
    response = requests.get(url, headers=headers)
    response.encoding = 'utf8'
    text = re.findall('<br />(.*)<br />', response.text, re.S)
    text = text[0].replace('&nbsp;', '').replace('<br />', '\n')
    print(f'{text[:100]}')
    return text


def save_to_word(t, text):
    """保存 word"""
    # 创建一个新的Word文档
    doc = Document()

    # 添加标题
    doc.add_heading(t, 0)  # 0代表最大的标题级别

    # 添加文本
    for line in text.split('\n'):
        doc.add_paragraph(line)

    # 保存Word文档
    doc.save(f'data/word/{t}.docx')


def save_to_pdf(t, content):
    """保存 pdf"""
    # 保存的 pdf 不会自动换行，会少字，处理一下
    new_content = ''
    for c in range(0, len(content), 43):
        new_content += content[c:c + 43] + '\n'

    # 注册中文字体
    pdfmetrics.registerFont(TTFont('msyh', 'data/msyh.ttc'))

    # 创建PDF文件
    pdf_filename = f'data/pdf/{t}.pdf'
    c = canvas.Canvas(pdf_filename, pagesize=letter)
    width, height = letter  # 获取页面宽度和高度
    width += 20

    # 添加标题
    c.setFont("msyh", 16)
    c.drawCentredString(width / 2, height - 60, t)

    # 添加自动换行的文本
    c.setFont("msyh", 12)
    text = c.beginText(40, height - 100)
    text.setFont("msyh", 12)
    text.textLines(new_content)
    text.setWordSpace(20)  # 设置字间距，以避免单词之间的断行
    text.setCharSpace(20)
    c.drawText(text)

    # 保存PDF
    c.save()


def save_to_txt(t, text):
    """保存 txt"""
    with open(f'data/txt/{t}.txt', 'w', encoding='utf-8') as file:
        file.write(t + '\n')
        file.write(text)


def save_to_md(t, text):
    """保存为 markdown"""
    with open(f'data/markdown/{t}.md', 'w', encoding='utf-8') as file:
        file.write('# ' + t + '\n')
        file.write(text)


if __name__ == '__main__':
    # 获得所有章节标题及 url
    titles, links = get_title_url()

    # 开始遍历爬取每一章节
    for t, u in list(zip(titles, links)):
        print(f'开始爬取 {t} 章节')
        try:
            text = get_text(u)

            # 随机保存到不同文件类型
            file_type = ['word', 'pdf', 'txt', 'md']
            type_ = random.choice(file_type)
            if type_ == 'word':
                save_to_word(t, text)
            elif type_ == 'pdf':
                save_to_pdf(t, text)
            elif type_ == 'txt':
                save_to_txt(t, text)
            else:
                save_to_md(t, text)

            print(f'保存成功!')
            print('========================================')
        except:
            pass
        time.sleep(1)
